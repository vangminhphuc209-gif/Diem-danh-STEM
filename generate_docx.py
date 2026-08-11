import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_hyperlink(paragraph, url, text, color="0056B3", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} {nsdecls("r")} r:id="{r_id}"/>')
    new_run = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:color w:val="{color}"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r>')
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def build_docx():
    doc = Document()
    
    # Page Setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x2B, 0x2D, 0x42)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("DANH SÁCH THIỆT BỊ LẮP RÁP CỤ THỂ")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("HỆ THỐNG ĐIỂM DANH THÔNG MINH STEM v4.0\n(QR Code + Nhận Diện Khuôn Mặt + Kiểm Tra Trang Phục AI)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    # Overview Box
    table_intro = doc.add_table(rows=1, cols=1)
    table_intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_intro = table_intro.cell(0, 0)
    set_cell_background(cell_intro, "F0F4F8")
    set_cell_margins(cell_intro, top=140, bottom=140, left=180, right=180)
    
    p_intro = cell_intro.paragraphs[0]
    p_intro.paragraph_format.space_after = Pt(0)
    r_intro = p_intro.add_run("📌 TỔNG QUAN TÀI LIỆU:\n")
    r_intro.bold = True
    r_intro.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    p_intro.add_run(
        "Báo cáo chi tiết này cung cấp danh sách đầy đủ toàn bộ phần cứng, linh kiện lắp ráp và phần mềm cần thiết "
        "để triển khai Hệ Thống Điểm Danh Thông Minh STEM v4.0. Mỗi thiết bị bao gồm tên cụ thể, giá thành tham khảo, "
        "đường dẫn mua hàng uy tín, khả năng liên kết hệ thống, thông số kỹ thuật chi tiết, hướng dẫn sử dụng và ghi chú dễ hiểu."
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 1: HARDWARE
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(8)
    r_h1 = p_h1.add_run("PHẦN I: DANH SÁCH THIẾT BỊ PHẦN CỨNG LẮP RÁP")
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    devices = [
        {
            "num": "1",
            "name": "Bo Mạch Máy Tính Mini Raspberry Pi 4 Model B (4GB RAM)",
            "category": "Lõi xử lý trung tâm (Brain/AI Controller)",
            "price": "1.500.000 – 1.650.000 VNĐ",
            "links": [
                ("NShop VN", "https://nshopvn.com/product/raspberry-pi-4-model-b-4gb/"),
                ("Raspberry Pi VN", "https://raspberrypi.vn/shop/raspberry-pi-4-model-b-4gb"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=raspberry%20pi%204%204gb")
            ],
            "connect": "• Cổng USB 3.0/2.0: Cắm Webcam USB.\n• Cổng Micro-HDMI: Nối ra Màn hình hiển thị HUD điểm danh.\n• Cổng USB-C: Cắm nguồn điện 5V/3A.\n• Khe MicroSD: Cắm thẻ nhớ lưu HĐH và dữ liệu dự án.\n• WiFi / LAN: Kết nối mạng Router để sync Google Sheets & truyền stream MJPEG.",
            "specs": "• CPU: Broadcom BCM2711, Quad-core Cortex-A72 (ARM v8) 64-bit SoC @ 1.5GHz.\n• RAM: 4GB LPDDR4-3200 SDRAM (Khuyến nghị 4GB/8GB để chạy mượt AI).\n• Kết nối không dây: 2.4 GHz và 5.0 GHz IEEE 802.11ac wireless, Bluetooth 5.0, BLE.\n• Mạng dây: Gigabit Ethernet.\n• Cổng giao tiếp: 2 cổng USB 3.0, 2 cổng USB 2.0, 2 cổng Micro-HDMI (hỗ trợ hiển thị đến 4K60).",
            "usage": "1. Dùng phần mềm Raspberry Pi Imager nạp HĐH Raspberry Pi OS (64-bit) vào thẻ MicroSD.\n2. Cắm thẻ nhớ, cắm nguồn USB-C, cắm camera USB và dây HDMI sang màn hình.\n3. Mở Terminal, clone mã nguồn dự án, cài đặt thư viện (`pip install -r requirements.txt`) và chạy `python main.py`.",
            "detail": "Đây là 'bộ não' trung tâm của cả hệ thống. Có kích thước nhỏ gọn chỉ bằng bàn tay (thẻ ATM) nhưng mang sức mạnh của một máy tính hoàn chỉnh, chịu trách nhiệm nhận luồng camera, chạy mô hình AI TFLite MobileNetV2 (nhận diện trang phục dân tộc), phân tích màu HSV (áo trắng, áo đoàn) và nhận diện khuôn mặt."
        },
        {
            "num": "2",
            "name": "Webcam USB Full HD 1080p (Logitech C920 / C920e / C922)",
            "category": "Cảm biến hình ảnh đầu vào (Vision Sensor)",
            "price": "900.000 – 1.250.000 VNĐ",
            "links": [
                ("CellphoneS", "https://cellphones.com.vn/webcam-logitech-c920e.html"),
                ("TNC Store", "https://www.tnc.com.vn/webcam-logitech-c920-pro-full-hd.html"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=webcam%20logitech%20c920")
            ],
            "connect": "• Cắm trực tiếp cổng cáp USB vào cổng USB 3.0 (màu xanh) hoặc USB 2.0 trên Raspberry Pi 4 hoặc PC Windows.\n• Thư viện OpenCV tự động quét và nhận dạng thiết bị qua tham số `CAMERA_INDEX = 0`.",
            "specs": "• Độ phân giải: Full HD 1080p (1920x1080) @ 30fps / HD 720p @ 30fps.\n• Tính năng lấy nét: Tự động lấy nét cao cấp (Autofocus sắc nét khi học sinh giơ thẻ QR sát camera).\n• Góc nhìn (FOV): 78° diagonal (rộng, bao quát được thân người học sinh đứng trước cửa).\n• Chiều dài dây USB: 1.5 mét.",
            "usage": "1. Gắn webcam lên chân Tripod hoặc đặt cố định trên bàn điểm danh / khung cửa ra vào ở độ cao 1.2m - 1.5m.\n2. Hướng góc quay thẳng vào luồng học sinh đi vào.\n3. Chạy lệnh `python step1_hsv_calibrator.py` để xem trực tiếp và hiệu chỉnh dải màu HSV cho phù hợp với ánh sáng phòng học.",
            "detail": "Là 'đôi mắt' của hệ thống. Độ nét 1080p cùng khả năng lấy nét tự động giúp quét mã QR cực nhanh trong vài miligiây và truyền hình ảnh chất lượng cao để AI phân tích khuôn mặt cũng như trang phục."
        },
        {
            "num": "3",
            "name": "IP Camera WiFi 1080p (Imou Ranger 2 / Hikvision DS-2CD) - Tùy chọn",
            "category": "Cảm biến hình ảnh không dây (Wireless Camera Option)",
            "price": "500.000 – 850.000 VNĐ",
            "links": [
                ("CellphoneS", "https://cellphones.com.vn/camera-ip-khong-day-imou-ranger-2-a22ep-l.html"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=camera%20imou%20ranger%202")
            ],
            "connect": "• Kết nối không dây WiFi 2.4GHz với Router.\n• Truyền luồng video trực tiếp RTSP qua mạng LAN đến Raspberry Pi / PC theo đường dẫn IP (Ví dụ: `rtsp://admin:password@192.168.1.50:554/live`).",
            "specs": "• Độ phân giải: 2.0 Megapixel Full HD 1080p.\n• Kết nối: WiFi IEEE802.11b/g/n hoặc cổng mạng dây RJ45.\n• Giao thức truyền tải: RTSP / ONVIF.\n• Nguồn cấp: Micro-USB 5V/1A.",
            "usage": "1. Kết nối Camera vào ứng dụng Imou/Hik-Connect trên điện thoại để cấp WiFi.\n2. Đặt IP cố định cho Camera trong trang quản trị Router.\n3. Trong file `config.py`, thay `CAMERA_INDEX = 0` bằng URL RTSP của camera.",
            "detail": "Phương án thay thế Webcam USB khi cần treo camera lên trần nhà, góc tường cao hoặc đặt cách xa máy tính xử lý mà không muốn đi dây USB dài vướng víu."
        },
        {
            "num": "4",
            "name": "Màn Hình Hiển Thị HDMI 21.5 inch / TV LED (VSP / E-Dra / LG)",
            "category": "Giao diện hiển thị trực quan (Display Unit)",
            "price": "1.200.000 – 2.200.000 VNĐ",
            "links": [
                ("CellphoneS", "https://cellphones.com.vn/man-hinh-vsp-21-5-inch-ve2152h.html"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=man%20hinh%2022%20inch%20hdmi")
            ],
            "connect": "• Nối từ cổng Micro-HDMI của Raspberry Pi (hoặc cổng HDMI của PC) tới cổng HDMI trên màn hình qua dây cáp Micro-HDMI to HDMI.",
            "specs": "• Kích thước: 21.5 inch – 23.8 inch.\n• Độ phân giải: Full HD 1920x1080 @ 60Hz/75Hz.\n• Cổng giao tiếp: HDMI, VGA.\n• Tấm nền: IPS / VA chống chói.",
            "usage": "1. Đặt màn hình tại bàn giám sát hoặc treo tường bên cạnh camera điểm danh.\n2. Khi hệ thống chạy `main.py`, màn hình hiển thị ngay lập tức luồng video camera kèm các ô màu trực quan: Xanh lá (Đã điểm danh/Trang phục OK), Đỏ (Chưa quét thẻ/Sai trang phục), Cam (Cần chụp ảnh xác minh UNCLEAR).",
            "detail": "Cung cấp phản hồi thị giác tức thì cho học sinh và bảo vệ/giáo viên biết trạng thái điểm danh và trang phục ngay tại chỗ."
        },
        {
            "num": "5",
            "name": "Thẻ Nhớ MicroSD SanDisk Extreme / Ultra 64GB Class 10 A2",
            "category": "Bộ nhớ lưu trữ chính (Storage Drive)",
            "price": "180.000 – 250.000 VNĐ",
            "links": [
                ("CellphoneS", "https://cellphones.com.vn/the-nho-microsd-sandisk-ultra-64gb-100mb-s.html"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=the%20nho%20sandisk%2064gb%20a2")
            ],
            "connect": "• Cắm trực tiếp vào khe thẻ nhớ MicroSD ở mặt dưới của Raspberry Pi 4.",
            "specs": "• Dung lượng: 64GB (thoải mái chứa HĐH, dữ liệu ảnh học sinh và mô hình AI).\n• Chuẩn tốc độ: A2 (Application Performance Class 2) - tối ưu đọc ghi ngẫu nhiên cho ứng dụng.\n• Tốc độ đọc: Lên đến 120MB/s - 160MB/s.\n• Chuẩn: UHS-I U3 V30 Class 10.",
            "usage": "1. Dùng đầu đọc thẻ nhớ cắm vào máy tính để flash HĐH Raspberry Pi OS 64-bit.\n2. Lưu trữ toàn bộ thư mục code `Diem-danh-STEM`, file cơ sở dữ liệu SQLite `database.db` và các mô hình AI `uniform_model_v2.tflite`.",
            "detail": "Là 'ổ cứng' chứa toàn bộ chương trình. Bắt buộc phải chọn thẻ nhớ đạt chuẩn A2 để đảm bảo tốc độ đọc ghi nhanh, tránh hiện tượng giật lag hoặc hỏng file hệ thống khi mất điện đột ngột."
        },
        {
            "num": "6",
            "name": "Thẻ QR Code Học Sinh (In Nhựa PVC 3 Lớp / Ép Plastic)",
            "category": "Thẻ định danh cá nhân (Student Identity Badge)",
            "price": "3.000 – 10.000 VNĐ / thẻ (~200.000 VNĐ / lớp 40 HS)",
            "links": [
                ("Xưởng In Thẻ PVC Shopee", "https://shopee.vn/search?keyword=in%20the%20nhua%20pvc%20theo%20yeu%20cau")
            ],
            "connect": "• Học sinh cầm/đeo thẻ trước camera. Thư viện `pyzbar` tự động nhận diện mã QR, trích xuất Mã Học Sinh (VD: `HS001`) và kích hoạt quá trình kiểm tra khuôn mặt + trang phục.",
            "specs": "• Kích thước: Chuẩn thẻ ATM CR80 (85.6 x 54 mm, dày 0.76mm).\n• Chất liệu: Nhựa PVC 3 lớp mờ/bóng chống nước, không phai màu.\n• Mã hóa: Mã QR chứa mã định danh học sinh duy nhất.",
            "usage": "1. Chạy lệnh `python generate_qr.py --pdf` để tạo sẵn file PDF thẻ QR cho toàn bộ học sinh trong danh sách `students.csv`.\n2. Gửi file in nhựa PVC hoặc in ra giấy cứng ép Plastic, đục lỗ dây đeo cho học sinh mang theo khi đến trường.",
            "detail": "Phương thức xác thực siêu nhanh và chính xác 100%. Giúp loại bỏ hoàn toàn rủi ro nhận diện nhầm giữa các học sinh có nét mặt giống nhau."
        },
        {
            "num": "7",
            "name": "Bộ Nguồn Adapter USB-C 5V / 3A (Hoặc Nguồn Raspberry Pi Chính Hãng 15.3W)",
            "category": "Bộ cấp nguồn điện (Power Supply Unit)",
            "price": "120.000 – 250.000 VNĐ",
            "links": [
                ("NShop VN", "https://nshopvn.com/product/nguon-raspberry-pi-4-5v-3a-cong-usb-c/"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=nguon%20raspberry%20pi%204%205v%203a")
            ],
            "connect": "• Đầu vào cắm điện lưới 220V AC. Đầu ra chân USB Type-C cắm vào cổng nguồn của Raspberry Pi 4.",
            "specs": "• Điện áp vào: 100-240V AC 50/60Hz.\n• Điện áp ra: 5.1V DC / 3.0A (15.3W).\n• Bảo vệ: Chống quá áp, quá dòng, ngắn mạch.",
            "usage": "1. Cắm chắc chắn vào cổng USB-C của Raspberry Pi trước khi cấp điện.\n2. Nên kết nối qua một cục Pin dự phòng (có hỗ trợ vừa sạc vừa xả) làm UPS mini chống cúp điện đột ngột.",
            "detail": "Cung cấp nguồn điện cho bo mạch. Rất quan trọng: Phải đảm bảo chuẩn dòng 3A để không bị lỗi 'Low Voltage' làm sụt hiệu năng chip AI."
        },
        {
            "num": "8",
            "name": "Bộ Router WiFi TP-Link Archer C20 / TL-WR841N / Router 4G",
            "category": "Thiết bị kết nối mạng (Network Connectivity)",
            "price": "300.000 – 550.000 VNĐ",
            "links": [
                ("CellphoneS", "https://cellphones.com.vn/bo-phat-wifi-tp-link-tl-wr841n.html"),
                ("Shopee Mall", "https://shopee.vn/search?keyword=router%20wifi%20tp-link")
            ],
            "connect": "• Phát mạng WiFi kết nối với Raspberry Pi.\n• Đẩy dữ liệu điểm danh trực tuyến lên Google Sheets Cloud qua Internet.\n• Tạo mạng nội bộ LAN cho phép điện thoại/máy tính khác mở trình duyệt xem camera trực tiếp tại `http://<IP-Raspberry>:8080`.",
            "specs": "• Băng thông: N300Mbps (2.4GHz) hoặc AC750 Dual-Band (2.4GHz & 5GHz).\n• Cổng kết nối: 1 WAN + 4 LAN Fast Ethernet.\n• Anten: 2 - 3 Anten ngoài 5dBi phủ sóng rộng.",
            "usage": "1. Cắm dây mạng Internet nhà trường vào cổng WAN Router.\n2. Kết nối Raspberry Pi vào WiFi do Router phát ra.",
            "detail": "'Cầu nối' thông tin giúp hệ thống vừa cập nhật dữ liệu điểm danh lên mây realtime, vừa cho phép thầy cô giám sát camera qua điện thoại."
        },
        {
            "num": "9",
            "name": "Máy Tính PC / Laptop Windows 10/11 64-bit (Tùy chọn Máy Chính / Training)",
            "category": "Máy chủ huấn luyện AI & Chạy thay thế (Workstation Server)",
            "price": "Thiết bị sẵn có (hoặc ~5.000.000 – 10.000.000 VNĐ nếu trang bị mới)",
            "links": [
                ("CellphoneS Laptop", "https://cellphones.com.vn/laptop.html"),
                ("Shopee Mall Laptop", "https://shopee.vn/search?keyword=laptop%20dell%20core%20i5")
            ],
            "connect": "• Cắm Webcam USB, kết nối WiFi/LAN. Có thể dùng thay thế hoàn toàn cho Raspberry Pi nếu không dùng máy nhúng.",
            "specs": "• Hệ điều hành: Windows 10 / 11 64-bit (BẮT BUỘC 64-BIT).\n• CPU: Intel Core i3/i5 thế hệ 8+ hoặc AMD Ryzen 3/5.\n• RAM: Tối thiểu 8GB DDR4.\n• Thư viện: Python 3.10+ 64-bit, TensorFlow, OpenCV, face_recognition, dlib.",
            "usage": "1. Dùng để chạy script chuẩn bị dữ liệu `step1_data_prep_v2.py` và huấn luyện mô hình AI MobileNetV2 `uniform_trainer.py`.\n2. Chạy ứng dụng điểm danh chính `python main.py` với tốc độ rất nhanh.",
            "detail": "Máy tính có cấu hình mạnh giúp quá trình huấn luyện AI chỉ mất vài phút. Sau khi train ra file `.tflite` nhẹ, có thể copy sang Raspberry Pi để chạy."
        }
    ]

    for dev in devices:
        # Heading for Device
        p_dev = doc.add_paragraph()
        p_dev.paragraph_format.space_before = Pt(14)
        p_dev.paragraph_format.space_after = Pt(4)
        r_dev_num = p_dev.add_run(f"{dev['num']}. {dev['name']}\n")
        r_dev_num.font.size = Pt(12)
        r_dev_num.font.bold = True
        r_dev_num.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        r_cat = p_dev.add_run(f"📂 Phân loại: {dev['category']}")
        r_cat.font.size = Pt(10)
        r_cat.font.italic = True
        r_cat.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        # Table for Specs & Info
        tbl = doc.add_table(rows=6, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        # Set column widths
        widths = [Inches(1.8), Inches(4.9)]
        
        rows_data = [
            ("💰 Giá thành ước tính", dev['price']),
            ("🛒 Link mua hàng uy tín", dev['links']),
            ("🔗 Khả năng liên kết", dev['connect']),
            ("⚙️ Thông số kỹ thuật", dev['specs']),
            ("💡 Hướng dẫn sử dụng", dev['usage']),
            ("📝 Chi tiết dễ hiểu", dev['detail'])
        ]

        for idx, (label, val) in enumerate(rows_data):
            row = tbl.rows[idx]
            cell_lbl = row.cells[0]
            cell_val = row.cells[1]

            cell_lbl.width = widths[0]
            cell_val.width = widths[1]

            set_cell_background(cell_lbl, "EDF2F7")
            set_cell_background(cell_val, "FAFAFA" if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
            set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)

            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_after = Pt(0)
            r_lbl = p_lbl.add_run(label)
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.space_after = Pt(0)

            if label == "🛒 Link mua hàng uy tín":
                for l_idx, (l_name, l_url) in enumerate(val):
                    if l_idx > 0:
                        p_val.add_run("  |  ")
                    p_val.add_run(f"👉 {l_name}: ")
                    add_hyperlink(p_val, l_url, l_url)
            else:
                r_val = p_val.add_run(val)
                r_val.font.size = Pt(9.5)
                r_val.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 2: SOFTWARE & LIBRARIES
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(18)
    p_h2.paragraph_format.space_after = Pt(8)
    r_h2 = p_h2.add_run("PHẦN II: DANH SÁCH PHẦN MỀM VÀ THƯ VIỆN LẬP TRÌNH")
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sw_list = [
        ("Python 3.10+ (64-bit)", "Ngôn ngữ lập trình chính", "Miễn phí", "https://www.python.org/downloads/", "Nền tảng chạy toàn bộ ứng dụng điểm danh và AI. Bắt buộc dùng phiên bản 64-bit."),
        ("OpenCV (opencv-python >= 4.8.0)", "Xử lý ảnh & HOG Person Detector", "Miễn phí", "pip install opencv-python opencv-contrib-python", "Đọc luồng camera realtime, phát hiện thân người (HOG), phân tích dải màu HSV và vẽ HUD."),
        ("TensorFlow / TFLite Runtime", "Chạy mô hình AI trang phục", "Miễn phí", "https://www.tensorflow.org/lite", "Chạy mô hình MobileNetV2 phân loại trang phục dân tộc (Thứ 2) với tốc độ ~30ms/frame trên CPU."),
        ("face-recognition & dlib", "Nhận diện khuôn mặt học sinh", "Miễn phí", "pip install cmake dlib face-recognition", "Nhận diện khuôn mặt học sinh chính xác cao, đối chiếu với cơ sở dữ liệu ảnh trong `data/stuface/`."),
        ("pyzbar & qrcode[pil]", "Quét & tạo mã QR Code", "Miễn phí", "pip install pyzbar qrcode[pil]", "pyzbar đọc mã QR từ camera < 1s. qrcode tạo thẻ QR cá nhân cho từng học sinh."),
        ("gspread & SQLite3", "Đồng bộ dữ liệu Cloud & Offline", "Miễn phí", "pip install gspread oauth2client", "gspread ghi điểm danh realtime lên Google Sheets. SQLite3 lưu dữ liệu dự phòng khi mất mạng."),
        ("Flask & Flask-CORS", "Web Streaming & REST API", "Miễn phí", "pip install flask flask-cors", "Phát luồng video camera MJPEG qua trình duyệt web tại đường dẫn `http://<IP>:8080`.")
    ]

    tbl_sw = doc.add_table(rows=1, cols=5)
    tbl_sw.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sw.autofit = False
    
    headers = ["Tên phần mềm / Thư viện", "Chức năng chính", "Chi phí", "Cú pháp cài / Link", "Mô tả sử dụng"]
    sw_widths = [Inches(1.5), Inches(1.4), Inches(0.8), Inches(1.5), Inches(1.5)]
    
    hdr_cells = tbl_sw.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = sw_widths[i]
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for item in sw_list:
        row = tbl_sw.add_row()
        for i in range(5):
            cell = row.cells[i]
            cell.width = sw_widths[i]
            set_cell_background(cell, "F7FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            val = item[i]
            if val.startswith("http"):
                add_hyperlink(p, val, "Tải về tại đây")
            else:
                r = p.add_run(val)
                r.font.size = Pt(8.5)
                if i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION 3: COST SUMMARY
    p_h3 = doc.add_paragraph()
    p_h3.paragraph_format.space_before = Pt(18)
    p_h3.paragraph_format.space_after = Pt(8)
    r_h3 = p_h3.add_run("PHẦN III: BẢNG TỔNG HỢP CHI PHÍ LẮP RÁP HỆ THỐNG")
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    cost_items = [
        ("1", "Raspberry Pi 4 Model B (4GB RAM)", "1", "1.500.000", "1.500.000 VNĐ", "Có thể dùng PC Windows sẵn có"),
        ("2", "Webcam USB Full HD 1080p Logitech C920", "1", "900.000", "900.000 VNĐ", "Khuyến nghị để QR sắc nét"),
        ("3", "Màn hình HDMI 21.5 inch Full HD", "1", "1.500.000", "1.500.000 VNĐ", "Có thể tận dụng TV/Màn cũ"),
        ("4", "Thẻ nhớ MicroSD SanDisk 64GB A2", "1", "220.000", "220.000 VNĐ", "Tối thiểu chuẩn Class 10 A1/A2"),
        ("5", "Adapter nguồn USB-C 5V/3A", "1", "120.000", "120.000 VNĐ", "Cung cấp nguồn cho Raspberry Pi"),
        ("6", "In Thẻ QR học sinh nhựa PVC (Lớp 40 HS)", "40", "5.000", "200.000 VNĐ", "40 thẻ x 5.000 VNĐ/thẻ"),
        ("7", "Router WiFi TP-Link N300", "1", "400.000", "400.000 VNĐ", "Có thể dùng WiFi sẵn của trường"),
        ("8", "Toàn bộ Phần mềm & Thư viện Python", "—", "0", "0 VNĐ", "Mã nguồn mở miễn phí 100%"),
        ("9", "Google Sheets Cloud API", "—", "0", "0 VNĐ", "Miễn phí theo tài khoản Google")
    ]

    tbl_cost = doc.add_table(rows=1, cols=6)
    tbl_cost.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cost.autofit = False

    cost_headers = ["STT", "Tên thiết bị / Hạng mục", "SL", "Đơn giá (VNĐ)", "Thành tiền (VNĐ)", "Ghi chú tối ưu chi phí"]
    c_widths = [Inches(0.5), Inches(2.2), Inches(0.5), Inches(1.1), Inches(1.2), Inches(1.2)]

    hdr_c = tbl_cost.rows[0].cells
    for i, title in enumerate(cost_headers):
        hdr_c[i].width = c_widths[i]
        set_cell_background(hdr_c[i], "2B6CB0")
        set_cell_margins(hdr_c[i], top=100, bottom=100, left=60, right=60)
        p = hdr_c[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for item in cost_items:
        row = tbl_cost.add_row()
        for i in range(6):
            cell = row.cells[i]
            cell.width = c_widths[i]
            set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=70, bottom=70, left=60, right=60)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in [3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            r = p.add_run(item[i])
            r.font.size = Pt(8.5)

    # Total row
    r_tot = tbl_cost.add_row()
    for i in range(6):
        cell = r_tot.cells[i]
        set_cell_background(cell, "EBF8FF")
        set_cell_margins(cell, top=90, bottom=90, left=60, right=60)

    cell_tot_lbl = r_tot.cells[0]
    p_t = cell_tot_lbl.paragraphs[0]
    r_t = p_t.add_run("TỔNG CHI PHÍ ƯỚC TÍNH LẮP RÁP MỚI HOÀN TOÀN:")
    r_t.font.bold = True
    r_t.font.size = Pt(9.5)
    r_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_val = r_tot.cells[4].paragraphs[0]
    p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_v = p_val.add_run("~4.840.000 VNĐ")
    r_v.font.bold = True
    r_v.font.size = Pt(10)
    r_v.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)

    # Save docx file
    output_dir = r"C:\Users\admin\Desktop\Diem-danh-STEM\baocao"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "DANH_SACH_THIET_BI_DIEM_DANH_STEM.docx")
    doc.save(output_path)
    print(f"File DOCX created successfully at: {output_path}")

if __name__ == "__main__":
    build_docx()
