import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os, sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Dark Navy
    elif level == 2:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Medium Blue
    else:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Dark Slate
    return p

def add_step_block(doc, step_num, step_title, target, tools_hw, tools_sw, tools_data, adv, diff_sol):
    p_st = doc.add_paragraph()
    p_st.paragraph_format.space_before = Pt(10)
    p_st.paragraph_format.space_after = Pt(4)
    r_st = p_st.add_run(f"📌 BƯỚC {step_num}: {step_title.upper()}")
    r_st.font.name = 'Arial'
    r_st.font.size = Pt(11.5)
    r_st.font.bold = True
    r_st.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # Box / Table for step breakdown
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    widths = [Inches(1.8), Inches(4.9)]

    rows_info = [
        ("🎯 Mục tiêu thực hiện", target),
        ("🛠️ Cần sử dụng những gì", f"• Phần cứng: {tools_hw}\n• Phần mềm & Thư viện: {tools_sw}\n• Dữ liệu / Cấu hình: {tools_data}"),
        ("✨ Thuận lợi trong bước này", adv),
        ("⚠️ Khó khăn gặp phải", diff_sol[0]),
        ("💡 Biện pháp khắc phục", diff_sol[1])
    ]

    for idx, (lbl, val) in enumerate(rows_info):
        row = tbl.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]

        cell_lbl.width = widths[0]
        cell_val.width = widths[1]

        set_cell_background(cell_lbl, "EDF2F7")
        set_cell_background(cell_val, "FAFAFA" if idx % 2 == 0 else "FFFFFF")
        set_cell_margins(cell_lbl, top=70, bottom=70, left=90, right=90)
        set_cell_margins(cell_val, top=70, bottom=70, left=90, right=90)

        p_l = cell_lbl.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(0)
        r_l = p_l.add_run(lbl)
        r_l.font.name = 'Arial'
        r_l.font.bold = True
        r_l.font.size = Pt(9.5)
        r_l.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

        p_v = cell_val.paragraphs[0]
        p_v.paragraph_format.space_after = Pt(0)
        r_v = p_v.add_run(val)
        r_v.font.name = 'Arial'
        r_v.font.size = Pt(9.5)
        r_v.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def append_detailed_execution_section(doc):
    add_styled_heading(doc, "QUY TRÌNH CÁC BƯỚC THỰC HIỆN THỰC TẾ & PHÂN TÍCH THUẬN LỢI - KHÓ KHĂN CHI TIẾT", level=1)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(8)
    r_desc = p_desc.add_run(
        "Để đưa hệ thống Điểm Danh Thông Minh STEM v4.0 từ ý tưởng lý thuyết vào vận hành thực tế tại Trường PTDTNT THPT tỉnh Lai Châu, "
        "nhóm nghiên cứu đã trải qua 8 bước thi công và thử nghiệm chặt chẽ. Mỗi bước đều được xác định rõ mục tiêu, trang thiết bị công cụ cần dùng, "
        "đồng thời đúc rút những thuận lợi cũng như biện pháp vượt qua các khó khăn kỹ thuật phát sinh."
    )
    r_desc.font.name = 'Arial'
    r_desc.font.size = Pt(10.5)

    steps_data = [
        {
            "num": "1",
            "title": "Khảo sát thực tế hiện trường lớp học & Lên sơ đồ kiến trúc hệ thống",
            "target": "Đánh giá luồng di chuyển của học sinh vào giờ truy bài, xác định vị trí đặt camera (khoảng cách 1m–2m, độ cao 1.3m–1.5m) và đo lường cường độ ánh sáng thực tế tại cửa lớp.",
            "hw": "Thước dây đo khoảng cách, máy đo độ sáng luxmeter/điện thoại, chân tripod thử nghiệm.",
            "sw": "Công cụ vẽ sơ đồ kiến trúc Draw.io / Visio, phần mềm quay thử camera.",
            "data": "Sơ đồ mặt bằng lớp học 11B1/12C1, thông số kích thước cửa ra vào.",
            "adv": "Được Ban Giám hiệu và Cô giáo hướng dẫn Nguyễn Thị Thúy Hạnh tạo điều kiện tối đa. Vị trí cửa lớp có luồng đi vào một chiều rõ ràng, rất thuận lợi cho việc cố định góc quay camera.",
            "diff": (
                "Ánh sáng môi trường tại cửa lớp biến đổi liên tục theo thời tiết (sáng sớm bị tối, trưa nắng chiếu chói hắt vào ống kính) khiến hình ảnh học sinh bị ngược sáng hoặc nhòe màu.",
                "Lắp đặt thêm cụm đèn LED trợ sáng nhỏ cố định trên khung camera để ổn định nguồn sáng, đồng thời thiết kế bộ công cụ hiệu chỉnh dải màu HSV linh hoạt bằng phần mềm."
            )
        },
        {
            "num": "2",
            "title": "Tập hợp dữ liệu học sinh, khởi tạo CSDL & Tạo thẻ QR Code định danh hàng loạt",
            "target": "Thu thập danh sách học sinh (Mã HS, Họ tên, Lớp), chụp ảnh chân dung làm dữ liệu nhận diện khuôn mặt và tự động xuất file thẻ QR Code chuẩn để in phát cho học sinh.",
            "hw": "Camera/Điện thoại HD chụp ảnh chân dung nền trắng, Máy in thẻ nhựa PVC / ép Plastic, Máy đục lỗ thẻ.",
            "sw": "Python `qrcode[pil]`, `Pillow`, `reportlab` (xuất PDF in), Microsoft Excel/CSV (`students.csv`).",
            "data": "Danh sách học sinh lớp 11B1 (`students.csv`), thư mục chứa ảnh chân dung chuẩn `data/stuface/`.",
            "adv": "Các bạn học sinh nhiệt tình phối hợp chụp ảnh mẫu. Script `generate_qr.py` chạy tự động hóa 100%, chỉ mất 10 giây để tạo toàn bộ thẻ QR dạng mã hóa nét cao và dàn trang sẵn file PDF sẵn sàng đem in.",
            "diff": (
                "Một số ảnh chân dung ban đầu chụp bị góc nghiêng, bóng đổ hoặc lộ ít khuôn mặt khiến thư viện nhận diện báo lỗi `IndexError` (không tìm thấy khuôn mặt trong ảnh).",
                "Ban hành quy chuẩn chụp ảnh chân dung (đứng thẳng góc 90 độ, đủ sáng, lộ trán và ears, 600x600px) và viết script `step1_dataset_audit.py` để tự động quét kiểm tra & lọc bỏ ảnh lỗi trước khi nạp vào CSDL."
            )
        },
        {
            "num": "3",
            "title": "Lập trình & Phát triển Module Nhận diện khuôn mặt kép (Dual-Layer Face Engine)",
            "target": "Xây dựng bộ thuật toán nhận diện khuôn mặt vừa đảm bảo độ chính xác >95%, vừa duy trì tốc độ mượt mà (20-25 FPS) trên thiết bị phần cứng phổ thông mà không bị lag.",
            "hw": "Máy tính PC / Laptop Windows 10/11 64-bit, Webcam Logitech C920.",
            "sw": "Python 3.10+ 64-bit, `opencv-python` (Haar Cascade), `dlib`, `face_recognition`, `numpy`, `cmake`.",
            "data": "File cấu hình mô hình Haar Cascade `haarcascade_frontalface_default.xml`, tập vector khuôn mặt 128 chiều.",
            "adv": "Thư viện `face_recognition` của Adam Geitgey dựa trên dlib ResNet có độ chính xác đỉnh cao (99.38%). Thuật toán Haar Cascade cực nhẹ, tốn gần như 0% CPU cho bước phát hiện vùng mặt.",
            "diff": (
                "1) Máy tính ban đầu chạy Python 32-bit (win32) nên bị lỗi không thể cài đặt được `dlib` và `face-recognition`.\n2) Khi chạy suy luận `face_recognition` liên tục ở 30 FPS, CPU bị đẩy lên 100%, làm khung hình bị giật khựng nghiêm trọng (sụt xuống 5-6 FPS).",
                "1) Nâng cấp toàn bộ môi trường sang Python 3.10 64-bit (amd64) kết hợp biên dịch C++ qua `cmake`.\n2) Thiết kế kiến trúc 'Dual-Layer': Dùng Haar Cascade phát hiện & vẽ khung mặt ở MỌI frame, nhưng chỉ kích hoạt `face_recognition` mã hóa 128 chiều ở MỖI 5 FRAME 1 LẦN (`FACE_SKIP_FRAMES = 5`). FPS tăng vọt lên 22-25 FPS mượt mà!"
            )
        },
        {
            "num": "4",
            "title": "Thiết kế & Huấn luyện Module AI Kiểm tra Trang phục Hybrid (MobileNetV2 TFLite + HSV)",
            "target": "Xây dựng bộ kiểm tra quy định trang phục tự động theo ngày: Thứ 2 (Trang phục dân tộc), Thứ 3&5 (Áo trắng + quần tối màu), Thứ 6 (Áo Đoàn TNCS + quần tối màu).",
            "hw": "Máy tính cá nhân GPU/CPU để train AI, bộ cảm biến HOG Person Detector.",
            "sw": "TensorFlow 2.13, TFLite Runtime, `scikit-learn`, `matplotlib`, `uniform_color_analyzer.py`.",
            "data": "Dataset `uniform_picdemo/` (~211 ảnh trang phục dân tộc, ảnh áo trắng, áo đoàn), file nhãn `uniform_labels_v2.json`.",
            "adv": "Áp dụng kỹ thuật Transfer Learning từ MobileNetV2 giúp thời gian huấn luyện cực ngắn (~5-10 phút). Xuất mô hình nén TFLite INT8 Quantization chỉ nhẹ 4.2 MB, chạy cực nhạy trên thiết bị nhúng.",
            "diff": (
                "1) Tập ảnh trang phục dân tộc ban đầu khá ít (~211 ảnh gốc), nguy cơ AI bị Quá khớp (Overfitting).\n2) Màu áo trắng và áo đoàn dễ bị nhận diện sai khi ánh sáng đèn điện phòng học phản chiếu.",
                "1) Áp dụng kỹ thuật Data Augmentation (`step1_data_prep_v2.py`) xoay, lật, zoom, chỉnh sáng nhân bản dữ liệu lên gấp 20 lần (>1.200 ảnh).\n2) Chuyển đổi màu từ RGB sang không gian HSV: Phân tích Áo trắng (Saturation thấp, Value cao) và Áo đoàn (Hue 95–135) chính xác tuyệt đối mà không phụ thuộc ánh sáng môi trường."
            )
        },
        {
            "num": "5",
            "title": "Xây dựng Kiến trúc Cơ sở Dữ liệu Kép (SQLite Offline Buffer + Google Sheets Cloud Sync)",
            "target": "Đảm bảo dữ liệu điểm danh được ghi nhận trực tuyến lên Google Sheets realtime cho giáo viên/phụ huynh tra cứu, đồng thời chống mất dữ liệu khi rớt mạng WiFi.",
            "hw": "Router WiFi TP-Link, cáp kết nối Internet, máy tính/Raspberry Pi.",
            "sw": "Google Cloud Console API v4, `gspread`, `oauth2client`, `sqlite3` Python, file JSON Service Account.",
            "data": "File CSDL SQLite `database.db` (bảng `students`, `logs`, `violations`), trang tính Google Sheets Online.",
            "adv": "Google Sheets API hoàn toàn miễn phí, giao diện bảng tính quen thuộc giúp giáo viên tra cứu & xuất file Excel dễ dàng. SQLite là CSDL nhúng cực nhẹ, không cần cài server daemon.",
            "diff": (
                "Mạng WiFi nhà trường có thời điểm bị chập chờn hoặc ngắt kết nối. Nếu ghi trực tiếp lên Google Sheets mà mất mạng, chương trình sẽ bị treo đơ 30 giây do chờ API timeout.",
                "Lập trình cơ chế Offline Buffer & Auto Sync: Điểm danh được ghi tức thì vào SQLite cục bộ trong 0.001s. Một luồng chạy ngầm (Background Thread) sẽ liên tục kiểm tra kết nối mạng: có mạng sẽ đẩy dần dữ liệu lên Cloud, mất mạng sẽ lưu an toàn trong SQLite và tự đồng bộ bù khi có mạng lại."
            )
        },
        {
            "num": "6",
            "title": "Thi công lắp ráp phần cứng, cân chỉnh vị trí Camera & Hiệu chuẩn dải màu HSV thực tế",
            "target": "Lắp ráp toàn bộ linh kiện điện tử thành thiết bị hoàn chỉnh, gá lắp cố định tại cửa lớp và chạy phần mềm cân chỉnh tham số màu sắc theo ánh sáng thực tế.",
            "hw": "Raspberry Pi 4 / Laptop, Webcam Logitech C920, Màn hình HDMI 21.5\", Cáp Micro-HDMI, Nguồn USB-C 5V/3A, Chân Tripod.",
            "sw": "Hệ điều hành Raspberry Pi OS / Windows, GUI Tool `step1_hsv_calibrator.py`.",
            "data": "File cấu hình dải màu `data/hsv_config.json`.",
            "adv": "Các thiết bị phần cứng chuẩn USB và HDMI đều hỗ trợ Plug-and-Play, lắp ráp nhanh chóng trong vòng 30 phút.",
            "diff": (
                "Ngưỡng màu HSV mặc định lập trình sẵn trong code bị lệch khi mang ra ánh sáng đèn tuýp của phòng học thực tế.",
                "Xây dựng phần mềm GUI tương tác `step1_hsv_calibrator.py`. Mở camera cho học sinh đứng trước ống kính, kéo các thanh trượt Trackbar trực quan để căn chỉnh dải màu HSV chính xác và nhấn phím `S` để tự động lưu vào `hsv_config.json`."
            )
        },
        {
            "num": "7",
            "title": "Tích hợp hệ thống chính (`main.py`), xây dựng Web Streaming & Dashboard báo cáo",
            "target": "Tích hợp toàn bộ các module riêng lẻ thành chương trình thống nhất `main.py`, phát luồng camera trực tiếp qua trang Web và tự động xuất báo cáo đồ họa HTML cuối ngày.",
            "hw": "Bộ thiết bị hoàn chỉnh, Điện thoại/Máy tính kết nối cùng mạng WiFi để xem stream.",
            "sw": "Python `Flask`, `Flask-CORS`, `threading`, `cv2`, `dashboard.py` (tích hợp đồ họa Chart.js).",
            "data": "Mẫu báo cáo HTML `bao_cao.html`, biểu đồ tròn & biểu đồ cột thống kê chuyên cần.",
            "adv": "Cấu trúc thiết kế dạng Modular giúp ghép nối các file python dễ dàng. Dashboard HTML đẹp mắt tự động tổng hợp số liệu Có mặt / Vắng mặt / Vi phạm.",
            "diff": (
                "1) Lỗi xung đột đa luồng (Thread-Lock): Luồng ghi SQLite và luồng Web Stream chạy cùng lúc gây khóa CSDL (`sqlite3.OperationalError: database is locked`).\n2) Rò rỉ RAM (RAM Leak) sau 4-5 tiếng chạy liên tục.",
                "1) Bổ sung đối tượng khóa luồng `threading.Lock()` bảo vệ mọi thao tác ghi CSDL.\n2) Tích hợp hàm dọn dẹp bộ nhớ (Garbage Collection) và tự động xóa bộ đệm mảng `recently_scanned` sau mỗi 300 khung hình."
            )
        },
        {
            "num": "8",
            "title": "Thử nghiệm thực tế tại lớp học (Lớp 11B1 & 12C1), đánh giá hiệu năng & Hoàn thiện",
            "target": "Vận hành hệ thống thực tế trong 2 tuần tại cửa lớp học, kiểm tra độ bền ổn định, thu thập đánh giá của giáo viên & học sinh để gỡ lỗi và hoàn thiện sản phẩm.",
            "hw": "Hệ thống điểm danh STEM v4.0 hoàn chỉnh đặt tại cửa lớp học.",
            "sw": "Chương trình `main.py`, hệ thống theo dõi Google Sheets, biểu mẫu khảo sát ý kiến.",
            "data": "Dữ liệu điểm danh thực tế 2 tuần của học sinh lớp 11B1 và 12C1.",
            "adv": "Học sinh rất hào hứng khi thấy tên và trạng thái trang phục 'OK' hiện màu xanh mượt mà trên màn hình. Giáo viên chủ nhiệm đánh giá cao vì không còn mất 10 phút điểm danh miệng.",
            "diff": (
                "Ngày Thứ 2 khi học sinh mặc trang phục dân tộc, có một số bạn mặc trang phục cách điệu mẫu mới khiến AI đưa ra kết quả `UNCLEAR` (không xác định được).",
                "Tận dụng tính năng tự động của hệ thống: Khi bị `UNCLEAR`, camera tự động chụp ảnh đẩy lên Cloud. Nhóm tải các ảnh này về, gắn nhãn bổ sung vào tập dữ liệu và chạy train lại mô hình AI. Độ chính xác nhận diện trang phục dân tộc tăng lên >95%!"
            )
        }
    ]

    for s in steps_data:
        add_step_block(
            doc, s['num'], s['title'], s['target'],
            s['hw'], s['sw'], s['data'], s['adv'], s['diff']
        )

    # TABLE SUMMARY OF ADVANTAGES & DIFFICULTIES
    add_styled_heading(doc, "BẢNG TỔNG HỢP PHÂN TÍCH THUẬN LỢI VÀ KHÓ KHĂN TRONG TOÀN BỘ QUÁ TRÌNH THỰC HIỆN", level=2)

    p_tbl_desc = doc.add_paragraph()
    p_tbl_desc.paragraph_format.space_after = Pt(6)
    r_td = p_tbl_desc.add_run(
        "Bảng dưới đây tổng hợp toàn bộ các yếu tố thuận lợi, khó khăn kỹ thuật phát sinh và các biện pháp giải quyết sáng tạo "
        "đã giúp nhóm nghiên cứu hoàn thiện sản phẩm đạt tiêu chuẩn chất lượng cao:"
    )
    r_td.font.name = 'Arial'
    r_td.font.size = Pt(10)

    tbl_sum = doc.add_table(rows=1, cols=4)
    tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sum.autofit = False

    t_headers = ["Hạng mục phân tích", "Các yếu tố Thuận lợi (Advantages)", "Những Khó khăn gặp phải (Difficulties)", "Biện pháp & Giải pháp khắc phục triệt để"]
    t_widths = [Inches(1.4), Inches(1.8), Inches(1.8), Inches(1.7)]

    hdr_cells = tbl_sum.rows[0].cells
    for i, title in enumerate(t_headers):
        hdr_cells[i].width = t_widths[i]
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], top=90, bottom=90, left=70, right=70)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    summary_rows = [
        (
            "1. Phần cứng & Môi trường",
            "• Được sự ủng hộ nhiệt tình của Ban Giám hiệu, Cô hướng dẫn và học sinh.\n• Tận dụng thiết bị có sẵn (Laptop, Camera điện thoại, TV cũ) tiết kiệm 90% chi phí.",
            "• Ánh sáng phòng học thay đổi thất thường theo thời tiết.\n• Bo mạch nhúng Raspberry Pi có tài nguyên CPU/RAM hạn chế.",
            "• Lắp đèn LED trợ sáng cố định & viết phần mềm GUI `step1_hsv_calibrator.py` chỉnh màu tại chỗ.\n• Tối ưu nén mô hình AI TFLite INT8 Quantization (~4.2MB)."
        ),
        (
            "2. Phần mềm & Thuật toán",
            "• Thư viện mã nguồn mở mạnh mẽ (OpenCV, TensorFlow, dlib, gspread).\n• Kiến trúc mã Python dạng Modular rất dễ mở rộng.",
            "• Đụng độ môi trường Python 32-bit không cài được `dlib`/`tensorflow`.\n• Thuật toán nhận diện mặt gốc bị sụt giảm FPS nghiêm trọng (5-6 FPS).",
            "• Nâng cấp hệ thống lên Python 3.10+ 64-bit hoàn chỉnh.\n• Sáng tạo thuật toán 2 lớp Dual-Layer (Haar Cascade tracking + face_recognition cách 5 frame) tăng vọt lên 25 FPS."
        ),
        (
            "3. Dữ liệu & Huấn luyện AI",
            "• Học sinh hào hứng tham gia chụp ảnh mẫu.\n• Học chuyển giao Transfer Learning MobileNetV2 giúp train AI chỉ mất 5-10 phút.",
            "• Tập ảnh trang phục dân tộc ban đầu ít (~211 ảnh), nguy cơ Overfitting.\n• Ảnh chân dung mẫu có ảnh bị nghiêng, mờ khiến dlib báo lỗi không thấy mặt.",
            "• Dùng Data Augmentation xoay, lật, zoom nhân bản tập dữ liệu gấp 20 lần (>1.200 ảnh).\n• Viết script `step1_dataset_audit.py` tự động quét lọc bỏ ảnh lỗi."
        ),
        (
            "4. Kết nối Mạng & CSDL",
            "• Google Sheets API miễn phí, dễ xem và xuất Excel.\n• SQLite là CSDL nhúng siêu nhẹ, không cần server daemon.",
            "• Mạng WiFi trường có lúc bị rớt làm ứng dụng bị treo đơ do chờ API timeout 30s.",
            "• Lập trình Kiến trúc CSDL Kép (SQLite Offline Buffer + Background Thread Sync tự đồng bộ khi có mạng lại)."
        ),
        (
            "5. Thử nghiệm Thực tế",
            "• Giao diện HUD màu sắc (Xanh/Đỏ/Cam) dễ quan sát.\n• Giảm thời gian điểm danh từ 10 phút xuống còn vài giây.",
            "• Mẫu trang phục dân tộc cách điệu mới lạ bị AI đánh giá nhầm thành `UNCLEAR`.",
            "• Tận dụng tính năng tự động chụp ảnh `UNCLEAR` upload Cloud → tải về trích xuất retraining bổ sung cho AI."
        )
    ]

    for item in summary_rows:
        row = tbl_sum.add_row()
        for i in range(4):
            cell = row.cells[i]
            cell.width = t_widths[i]
            set_cell_background(cell, "F7FAFC" if i == 0 else "FFFFFF")
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(item[i])
            r.font.name = 'Arial'
            r.font.size = Pt(8.5)
            if i == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

def update_file(file_path):
    print(f"Processing: {file_path}")
    doc = Document(file_path)
    append_detailed_execution_section(doc)
    doc.save(file_path)
    print(f"Successfully updated: {file_path}")

if __name__ == "__main__":
    folder = r"C:\Users\admin\Desktop\Diem-danh-STEM\baocao"
    files = [
        "Bao_Cao_KHKT_Diem_Danh_Thong_Minh.docx",
        "BAN THUYET MINH SAN PHAM HE THONG DIEM DANH THONG MINH.docx",
        "BAN_THUYET_MINH_SAN_PHAM.docx"
    ]
    for f in files:
        full_path = os.path.join(folder, f)
        if os.path.exists(full_path):
            update_file(full_path)
